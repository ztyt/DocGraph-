import json
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document

from docgraph_sidecar.core.db import connect, initialize_database
from docgraph_sidecar.parser import (
    DocxParser,
    ParseContext,
    ParserError,
    default_parser_registry,
    parse_with_error_recording,
)


class DocxParserTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.parser = DocxParser()

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_docx_parser_extracts_headings_paragraphs_and_tables(self) -> None:
        path = self.root / "contract.docx"
        document = Document()
        document.add_heading("Alpha Project", level=1)
        document.add_paragraph("Contract summary paragraph.")
        document.add_heading("Budget", level=2)
        document.add_paragraph("Owner: North Center.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Quantity"
        table.cell(1, 0).text = "Camera"
        table.cell(1, 1).text = "12"
        document.save(path)

        result = self.parser.parse(ParseContext.from_path(path, file_id="file-docx"))

        self.assertEqual(result.parser_name, "docx")
        self.assertEqual(
            [element.element_type for element in result.elements],
            ["heading", "paragraph", "heading", "paragraph", "table"],
        )
        self.assertEqual(result.elements[0].text, "Alpha Project")
        self.assertEqual(result.elements[2].section_path, "Alpha Project > Budget")
        self.assertEqual(result.elements[4].section_path, "Alpha Project > Budget")
        self.assertEqual(result.elements[4].metadata, {"rows": 2, "columns": 2})
        self.assertIn("Item | Quantity", result.chunks[4].text)
        self.assertEqual(result.chunks[4].heading, "Budget")

    def test_basic_fixture_docx_can_be_parsed(self) -> None:
        path = Path("fixtures/basic_docs/alpha_contract.docx")

        result = self.parser.parse(ParseContext.from_path(path, file_id="file-fixture"))

        self.assertEqual(len(result.elements), 1)
        self.assertEqual(result.elements[0].element_type, "paragraph")
        self.assertIn("Alpha Project contract summary", result.chunks[0].text)

    def test_default_registry_selects_docx_parser(self) -> None:
        path = self.root / "simple.docx"
        document = Document()
        document.add_paragraph("Simple body")
        document.save(path)

        result = default_parser_registry().parse_path(path, file_id="file-registry")

        self.assertEqual(result.parser_name, "docx")
        self.assertEqual(result.chunks[0].text, "Simple body")

    def test_bad_docx_records_parse_error(self) -> None:
        data_dir = self.root / "data"
        bad_path = self.root / "bad.docx"
        bad_path.write_bytes(b"not a zip package")
        initialize_database(data_dir=data_dir)
        self._insert_file(data_dir, "file-bad", str(bad_path))
        context = ParseContext.from_path(bad_path, file_id="file-bad")

        with self.assertRaises(ParserError):
            parse_with_error_recording(default_parser_registry(), context, data_dir=data_dir)

        connection = connect(data_dir=data_dir)
        try:
            row = connection.execute(
                """
                SELECT error_code, error_message, parser_name, retryable, details_json
                FROM parse_errors
                WHERE file_id = 'file-bad'
                """
            ).fetchone()
        finally:
            connection.close()

        self.assertIsNotNone(row)
        self.assertEqual(row["error_code"], "DOCX_PARSE_ERROR")
        self.assertEqual(row["parser_name"], "docx")
        self.assertEqual(row["retryable"], 0)
        self.assertEqual(row["error_message"], "DOCX file could not be parsed.")
        self.assertEqual(json.loads(row["details_json"])["path"], str(bad_path))

    def _insert_file(self, data_dir: Path, file_id: str, path: str) -> None:
        connection = connect(data_dir=data_dir)
        try:
            connection.execute(
                """
                INSERT INTO files (
                  file_id,
                  path,
                  normalized_path,
                  filename,
                  extension,
                  source_type,
                  file_status,
                  parse_status,
                  deleted_flag
                )
                VALUES (?, ?, ?, ?, '.docx', 'office', 'discovered', 'pending', 0)
                """,
                (file_id, path, path.casefold(), Path(path).name),
            )
            connection.commit()
        finally:
            connection.close()


if __name__ == "__main__":
    unittest.main()
